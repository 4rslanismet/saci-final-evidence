from __future__ import annotations

import argparse
import hashlib
import json
import zipfile
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn


def twips_to_inches(value: int | None) -> float | None:
    return None if value is None else round(value / 1440, 4)


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("docx", type=Path)
    parser.add_argument("--out", type=Path, required=True)
    args = parser.parse_args()

    document = Document(args.docx)
    paragraphs = []
    for index, paragraph in enumerate(document.paragraphs):
        runs = []
        for run in paragraph.runs:
            runs.append(
                {
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font": run.font.name,
                    "size_pt": None if run.font.size is None else run.font.size.pt,
                }
            )
        paragraphs.append(
            {
                "index": index,
                "style": paragraph.style.name if paragraph.style else None,
                "text": paragraph.text,
                "alignment": None if paragraph.alignment is None else int(paragraph.alignment),
                "format": {
                    "space_before_pt": None if paragraph.paragraph_format.space_before is None else paragraph.paragraph_format.space_before.pt,
                    "space_after_pt": None if paragraph.paragraph_format.space_after is None else paragraph.paragraph_format.space_after.pt,
                    "line_spacing": paragraph.paragraph_format.line_spacing,
                    "left_indent_in": None if paragraph.paragraph_format.left_indent is None else paragraph.paragraph_format.left_indent.inches,
                    "right_indent_in": None if paragraph.paragraph_format.right_indent is None else paragraph.paragraph_format.right_indent.inches,
                    "first_line_indent_in": None if paragraph.paragraph_format.first_line_indent is None else paragraph.paragraph_format.first_line_indent.inches,
                    "keep_with_next": paragraph.paragraph_format.keep_with_next,
                    "keep_together": paragraph.paragraph_format.keep_together,
                },
                "runs": runs,
            }
        )

    sections = []
    for index, section in enumerate(document.sections):
        sect_pr = section._sectPr
        cols = sect_pr.find(qn("w:cols"))
        sections.append(
            {
                "index": index,
                "start_type": int(section.start_type),
                "page_width_in": twips_to_inches(section.page_width.twips),
                "page_height_in": twips_to_inches(section.page_height.twips),
                "left_margin_in": twips_to_inches(section.left_margin.twips),
                "right_margin_in": twips_to_inches(section.right_margin.twips),
                "top_margin_in": twips_to_inches(section.top_margin.twips),
                "bottom_margin_in": twips_to_inches(section.bottom_margin.twips),
                "column_count": int(cols.get(qn("w:num"), "1")) if cols is not None else 1,
                "column_spacing_twips": int(cols.get(qn("w:space"), "0")) if cols is not None else None,
            }
        )

    used_style_names = sorted({item["style"] for item in paragraphs if item["style"]})
    styles = []
    for style_name in used_style_names:
        style = document.styles[style_name]
        paragraph_format = style.paragraph_format
        styles.append(
            {
                "name": style_name,
                "type": int(style.type),
                "font_name": style.font.name,
                "font_size_pt": None if style.font.size is None else style.font.size.pt,
                "bold": style.font.bold,
                "italic": style.font.italic,
                "alignment": None if paragraph_format.alignment is None else int(paragraph_format.alignment),
                "space_before_pt": None if paragraph_format.space_before is None else paragraph_format.space_before.pt,
                "space_after_pt": None if paragraph_format.space_after is None else paragraph_format.space_after.pt,
                "line_spacing": paragraph_format.line_spacing,
                "left_indent_in": None if paragraph_format.left_indent is None else paragraph_format.left_indent.inches,
                "right_indent_in": None if paragraph_format.right_indent is None else paragraph_format.right_indent.inches,
                "first_line_indent_in": None if paragraph_format.first_line_indent is None else paragraph_format.first_line_indent.inches,
                "keep_with_next": paragraph_format.keep_with_next,
                "keep_together": paragraph_format.keep_together,
            }
        )

    package_parts = []
    with zipfile.ZipFile(args.docx) as archive:
        for name in sorted(archive.namelist()):
            payload = archive.read(name)
            package_parts.append(
                {
                    "path": name,
                    "size": len(payload),
                    "sha256": hashlib.sha256(payload).hexdigest(),
                }
            )

    payload = {
        "path": str(args.docx.resolve()),
        "size": args.docx.stat().st_size,
        "sha256": hashlib.sha256(args.docx.read_bytes()).hexdigest(),
        "paragraph_count": len(paragraphs),
        "table_count": len(document.tables),
        "inline_shape_count": len(document.inline_shapes),
        "sections": sections,
        "styles": styles,
        "paragraphs": paragraphs,
        "package_parts": package_parts,
    }
    args.out.parent.mkdir(parents=True, exist_ok=True)
    args.out.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {args.out} with {len(paragraphs)} paragraphs")


if __name__ == "__main__":
    main()
